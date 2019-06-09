import datetime
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

document = Document('derwentwater_template.docx')

obj_styles = document.styles
obj_reportStyle = obj_styles.add_style('ReportStyle', WD_STYLE_TYPE.CHARACTER)
obj_report_font = obj_reportStyle.font
obj_report_font.size = Pt(9.5)
obj_report_font.name = 'Calibri'

obj_english_reportStyle = obj_styles.add_style('EnglishReportStyle', WD_STYLE_TYPE.CHARACTER)
obj_english_font = obj_english_reportStyle.font
obj_english_font.size = Pt(9.5)
obj_english_font.name = 'Calibri'


obj_dateStyle = obj_styles.add_style('DateStyle', WD_STYLE_TYPE.CHARACTER)
obj_font = obj_dateStyle.font
obj_font.size = Pt(14)
obj_font.name = 'Calibri'

obj_nameStyle = obj_styles.add_style('NameStyle', WD_STYLE_TYPE.CHARACTER)
obj_name_font = obj_nameStyle.font
obj_name_font.size = Pt(12)
obj_name_font.name = 'Calibri'
obj_name_font.bold = True

if len(sys.argv) < 6:
    print("ERROR: Not enough information. Please provide First Name, Last Name, Gender - f for female, m for male,"
          " Reading, Writing, Maths"
          "- where 1 is working towards, 2 is working at, 3 is working at greater depth.")
    exit()

firstName = sys.argv[1]
lastName = sys.argv[2]
fullName = firstName + " " + lastName
gender = sys.argv[3]

reading = int(sys.argv[4])
    
if reading < 1 or reading > 3:
    print("ERROR: Please insert a valid value for Reading Attainment - "
          "where 1 is working towards, 2 is working at, 3 is working at greater depth.")
    exit()

writing = int(sys.argv[5])
    
if writing < 1 or writing > 3:
    print("ERROR: Please insert a valid value for Writing Attainment - "
          "where 1 is working towards, 2 is working at, 3 is working at greater depth.")
    exit()

maths = int(sys.argv[6])

if maths < 1 or maths > 3:
    print("ERROR: Please insert a valid value for Maths Attainment - "
          "where 1 is working towards, 2 is working at, 3 is working at greater depth.")
    exit()

nominalPronoun = "she"
possessivePronoun = "her"
pointingPronoun = "her"
nominalPronounCapitalised = "She"
possessivePronounCapitalised = "Her"
pointingPronounCapitalised = "Her"
report = ""
readingReport = ""
writingReport = ""
mathsReport = ""
englishReport = ""

if gender != "m" and gender != "f":
    print("Sorry, this gender is not yet supported. Please enter f for female or m for male")
    exit()
elif gender == "m":
    nominalPronoun = "he"
    possessivePronoun = "his"
    pointingPronoun = "him"
    nominalPronounCapitalised = "He"
    possessivePronounCapitalised = "His"
    pointingPronounCapitalised = "Him"

levels = {
    1: "towards the expected level",
    2: "at the expected level",
    3: "at greater depth"
}

readingLevel = levels[reading]
writingLevel = levels[writing]
mathsLevel = levels[maths]

frontPageComments = {
    'reading': f"{firstName} is working {readingLevel} for Year 3 reading",
    'writing': f"{firstName} is working {writingLevel} for Year 3 writing",
    'maths': f"{firstName} is working {mathsLevel} for Year 3 mathematics",
}

keys = {
    3: 'M',
    2: 'M',
    1: 'W'
}

goodReading = (
    f"{firstName} reads with fluency and expression, and is confident about reading aloud to the rest of the class. "
    f"{nominalPronounCapitalised} carefully selects evidence from texts to support {possessivePronoun} answers during "
    f"guided reading, and makes erudite and insightful inferences about stories."
)

okReading = (
    f"{firstName} has worked very hard to achieve {possessivePronoun} targets in English this year and I am pleased "
    f"with {possessivePronoun} progress. "
    f"{nominalPronounCapitalised} is a confident reader, and uses evidence to support {possessivePronoun} answers "
    f"about texts. "
    f"{nominalPronounCapitalised} makes good predictions about texts, and {possessivePronoun} inference skills have "
    f"improved considerably. "
)

badReading = (
    f"{firstName} has worked hard this year to improve {possessivePronoun} reading and writing. "
    f"{nominalPronounCapitalised} reads with increasing confidence and accuracy, and is able to retrieve "
    f"information from a text. "
    f"{nominalPronounCapitalised} can explain the meaning of texts, and now needs to work on improving "
    f"{possessivePronoun} inference skills. "
)

goodWriting = (
    f""
)

okWriting = (
    f"{firstName} is a confident and imaginative writer, and the presentation of {possessivePronoun} work is good. "
    f"In order to further develop {possessivePronoun} writing skills, {firstName} should make sure to carefully read "
    f"through {possessivePronoun} work, to ensure {possessivePronoun} punctuation and spelling are always accurate. "
    f"I would also like {pointingPronoun} to continue to develop {possessivePronoun} use of interesting vocabulary "
    f"and subordinating clauses, in order to take {possessivePronoun} writing to the next level. "
)

badWriting = (
    f"{firstName} has also become an increasingly confident and independent writer as the year has progressed. "
    f"The presentation of {possessivePronoun} work is good, and {nominalPronoun} now writes a good amount each lesson. "
    f"{possessivePronounCapitalised} use of full stops and capital letters is sound, but not always consistent, and I "
    f"would like {pointingPronoun} to partition {possessivePronoun} sentences using more commas and full stops. "
    f"{possessivePronounCapitalised} spelling has improved steadily, and the clarity of {possessivePronoun} sentences "
    f"is good. "
)

goodMaths = (
    f"{firstName} is a very gifted mathematician and has worked extremely hard this year. "
    f"{nominalPronounCapitalised} is able to soundly recall all of {possessivePronoun} multiplication and division "
    f"facts and can apply these to a range of mathematical situations and word problems. "
    f"{firstName} works quickly and efficiently, taking the time to show {possessivePronoun} working out and to "
    f"explain {possessivePronoun} answers. "
    f"{nominalPronounCapitalised} can confidently and accurately use a range of written and mental methods to solve "
    f"complex multi-step problems. "
    f"{nominalPronounCapitalised} continues to challenge and stretch {pointingPronoun}self every day, and this has "
    f"helped {pointingPronoun} to make outstanding progress. "
)

okMaths = (
    f"{firstName} has made good progress in maths this year. "
    f"{nominalPronounCapitalised} can now confidently recall most of {possessivePronoun} times tables and division "
    f"facts, and can accurately apply these facts to solve word problems. "
    f"{firstName} can use a range of mental strategies to solve calculations and is becoming more confident "
    f"about using written methods for addition and subtraction. "
    f"{nominalPronounCapitalised} completes {possessivePronoun} work independently and sensibly at {possessivePronoun} "
    f"table, and presents {possessivePronoun} work neatly in {possessivePronoun} book. "
    f"{firstName} is confident about asking for help when {nominalPronoun} needs it, and this has helped "
    f"{pointingPronoun} to make good progress with {possessivePronoun} learning. "
)

badMaths = (
    f"I am pleased with {firstName}’s progress in maths this year. {nominalPronounCapitalised} can recall "
    f"{possessivePronoun} 3, 4 and 8 times tables and division facts with increasing confidence, and is starting "
    f"to apply these facts to solve word problems. {nominalPronounCapitalised} completes {possessivePronoun} work "
    f"independently and sensibly at {possessivePronoun} table, and presents {possessivePronoun} work neatly in "
    f"{possessivePronoun} book. {firstName} is confident about asking for help when {nominalPronoun} needs it, "
    f"and this has helped {pointingPronoun} make good progress with {possessivePronoun} learning. "
    f"{nominalPronounCapitalised} sometimes finds solving calculations in {possessivePronoun} head difficult, "
    f"but is developing a range of strategies to help {pointingPronoun} with this. "
)

report += "\nENGLISH: \n"
if reading == 1:
    readingReport += badReading
elif reading == 2:
    readingReport += okReading
else:
    readingReport += goodReading

if writing == 1:
    writingReport += badWriting
elif writing == 2:
    writingReport += okWriting
else:
    writingReport += goodWriting

englishReport += readingReport + writingReport
report += englishReport + "\n"

report += "\nMATHS: \n"
if maths == 1:
    mathsReport += badMaths
elif maths == 2:
    mathsReport += okMaths
else:
    mathsReport += goodMaths

report += mathsReport + "\n"


def addDetails():
    table = document.tables[0]
    first_column = table.columns[0]
    second_column = table.columns[1]

    first_column_paragraph = first_column.cells[0].paragraphs[0]

    if 'Name:' in first_column_paragraph.text:
        first_column_paragraph.add_run(" " + fullName, style = 'NameStyle')

    if 'Attainment' in second_column.cells[0].paragraphs[0].text:
        second_column.cells[1].paragraphs[0].add_run(
            frontPageComments['reading'] + '\n' +
            frontPageComments['writing'] + '\n' +
            frontPageComments['maths']
        )

    for para in document.paragraphs:
        if 'Date:' in para.text:
            para.add_run(f" {datetime.datetime.now():%d-%m-%Y}")
    return


def addSubjectAchieved(subject, subject_key):
    section_is_writing = False
    dict_of_end_words = {
        'reading': 'Writing',
        'writing': 'Comment',
        'maths': 'Comment'
    }
    table = document.tables[2]
    if subject == 'maths':
        table = document.tables[3]
    lstRows = iter(table.rows)
    for rows in lstRows:
        for cell in rows.cells:
            lst_paras = iter(cell.paragraphs)
            for paragraph in lst_paras:
                if paragraph.text == 'Writing':
                    section_is_writing = True
                if len(paragraph.text) < 1:
                    if subject != 'writing' and not section_is_writing:
                        paragraph.add_run(subject_key, style='ReportStyle')
                    if subject == 'writing' and section_is_writing:
                        paragraph.add_run(subject_key, style='ReportStyle')
                elif dict_of_end_words[subject]  == paragraph.text:
                    return


def addSubjectReport(subject, subject_report):
    table = document.tables[2]
    report_style = "EnglishReportStyle"
    if subject == 'maths':
        table = document.tables[3]
        report_style = "ReportStyle"
    lstRows = iter(table.rows)
    for rows in lstRows:
        for cell in rows.cells:
            lst_paras = iter(cell.paragraphs)
            for paragraph in lst_paras:
                if 'Comment' == paragraph.text:
                    for i_row in lstRows:
                        i_row.cells[1].paragraphs[0].add_run(subject_report, style=report_style)
                        return


addDetails()
addSubjectAchieved('writing', keys[writing])
addSubjectAchieved('reading', keys[reading])
addSubjectAchieved('maths', keys[maths])
addSubjectReport('english', englishReport)
addSubjectReport('maths', mathsReport)

documentName = firstName + '_' + lastName + '.docx'
document.save(documentName)

print(report)
