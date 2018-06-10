import sys
import random
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

document = Document('master.docx')

obj_styles = document.styles
obj_reportStyle = obj_styles.add_style('ReportStyle', WD_STYLE_TYPE.CHARACTER)
obj_font = obj_reportStyle.font
obj_font.size = Pt(10)
obj_font.name = 'Calibri'

obj_nameStyle = obj_styles.add_style('NameStyle', WD_STYLE_TYPE.CHARACTER)
obj_name_font = obj_nameStyle.font
obj_name_font.name = 'Calibri'

if len(sys.argv) < 9:
    print("ERROR: Not enough information. Please provide First Name, Last Name, Gender - f for female, m for male,"
          " RE, English, Phonics, Maths, Science"
          "- where 1 is bad, 2 is ok, 3 is good attainment or behaviour.")
    exit()

firstName = sys.argv[1]
lastName = sys.argv[2]
fullName = firstName + " " + lastName
gender = sys.argv[3]

religStudy = int(sys.argv[4][0])

if religStudy < 1 or religStudy > 3:
    print("ERROR: Please insert a valid value for Religious Study Attainment - where 1 is bad, 2 is ok,"
          " 3 is good Attainment.")
    exit()

try:
    religStudyEffort = int(sys.argv[4][1])
except IndexError:
    print("ERROR: Please insert a value for Religious Study Effort after the value for attainment, e.g. "
          "32 means an Ab for Religious Study attainment, because the first number is a 3, and an At for "
          "Religious Study Effort, because the second number is a 2.")
    exit()

if religStudyEffort < 1 or religStudyEffort > 3:
    print("ERROR: Please insert a valid value for Religious Study Effort - where 1 is bad, 2 is ok,"
          " 3 is good Attainment.")
    exit()

english = int(sys.argv[5][0])
    
if english < 1 or english > 3:
    print("ERROR: Please insert a valid value for English Attainment - where 1 is bad, 2 is ok, 3 is good Attainment.")
    exit()

try:
    englishEffort = int(sys.argv[5][1])
except IndexError:
    print("ERROR: Please insert a value for English Effort after the value for attainment, e.g. "
          "32 means an Ab for English attainment, because the first number is a 3, and an At for "
          "English Effort, because the second number is a 2.")
    exit()

if englishEffort < 1 or englishEffort > 3:
    print("ERROR: Please insert a valid value for English Effort - where 1 is bad, 2 is ok,"
          " 3 is good Attainment.")
    exit()

phonics = int(sys.argv[6][0])
    
if phonics < 1 or phonics > 3:
    print("ERROR: Please insert a valid value for Phonics Attainment - where 1 is bad, 2 is ok, 3 is good Attainment.")
    exit()

try:
    phonicsEffort = int(sys.argv[6][1])
except IndexError:
    print("ERROR: Please insert a value for Phonics Effort after the value for attainment, e.g. "
          "32 means an Ab for Phonics attainment, because the first number is a 3, and an At for "
          "Phonics Effort, because the second number is a 2.")
    exit()

if phonicsEffort < 1 or phonicsEffort > 3:
    print("ERROR: Please insert a valid value for Phonics Effort - where 1 is bad, 2 is ok,"
          " 3 is good Attainment.")
    exit()

maths = int(sys.argv[7][0])

if maths < 1 or maths > 3:
    print("ERROR: Please insert a valid value for Maths Attainment - where 1 is bad, 2 is ok, 3 is good Attainment.")
    exit()

try:
    mathsEffort = int(sys.argv[7][1])
except IndexError:
    print("ERROR: Please insert a value for Maths Effort after the value for attainment, e.g. "
          "32 means an Ab for Maths attainment, because the first number is a 3, and an At for "
          "Maths Effort, because the second number is a 2.")
    exit()

if mathsEffort < 1 or mathsEffort > 3:
    print("ERROR: Please insert a valid value for Maths Effort - where 1 is bad, 2 is ok,"
          " 3 is good Attainment.")
    exit()

science = int(sys.argv[8][0])
    
if science < 1 or science > 3:
    print("ERROR: Please insert a valid value for Science Attainment - where 1 is bad, 2 is ok, 3 is good Attainment.")
    exit()

try:
    scienceEffort = int(sys.argv[8][1])
except IndexError:
    print("ERROR: Please insert a value for Science Effort after the value for attainment, e.g. "
          "32 means an Ab for Science attainment, because the first number is a 3, and an At for "
          "Science Effort, because the second number is a 2.")
    exit()

if scienceEffort < 1 or scienceEffort > 3:
    print("ERROR: Please insert a valid value for Science Effort - where 1 is bad, 2 is ok,"
          " 3 is good Attainment.")
    exit()

nominalPronoun = "she"
possessivePronoun = "her"
pointingPronoun = "her"
nominalPronounCapitalised = "She"
possessivePronounCapitalised = "Her"
pointingPronounCapitalised = "Her"
report = ""
religStudyReport = ""
englishReport = ""
mathsReport = ""
scienceReport = ""
generalReport = ""

if gender != "m" and gender != "f":
    print("Sorry, this gender is not yet supported. Please enter f for female or m for male")
    exit()
elif gender == "m":
    nominalPronoun = "he"
    possessivePronoun = "his"
    pointingPronoun = "him"
    nominalPronounCapitalised = "He"
    possessivePronounCapitalised = "His"
    pointingPronounCapitalised = "Him"

goodMathsStart = (
    "{0} has a sound and in-depth understanding of a range of mathematical concepts, and I have been very "
    "impressed by {2} continued progress in maths this year. ",
    "{0}’s progress in maths this year has been excellent. "
)

goodMathsMiddle = (
    "{0} is confident with using a range of strategies to solve two-digit addition and subtraction problems, "
    "and is able to use mathematical reasoning to prove {2} answers both verbally and in writing. {3} takes "
    "great care with the presentation of {2} sums in {2} maths book. ",
    "{3} demonstrates a good awareness of a range of strategies which can be used to solve addition and subtraction "
    "problems, and is able to explain {2} answers clearly both verbally and in writing. {3} can solve equations "
    "involving two-digit numbers, and presents {2} working out neatly in {2} maths book. "
)

goodMathsEnd = (
    "{3} consistently demonstrates a comprehensive understanding of the link between abstract and concrete "
    "representations of equations, and can confidently use the Dienes blocks and string-beads to prove {2} answers. ",
    "{3} is able to use concrete resources, such as the Dienes blocks and string-beads to prove {2} answers, and "
    "can confidently link these representations to abstract equations. "
)

goodMaths = (
    goodMathsStart[0] + goodMathsMiddle[0] + goodMathsEnd[0],
    goodMathsStart[0] + goodMathsMiddle[1] + goodMathsEnd[0],
    goodMathsStart[0] + goodMathsMiddle[0] + goodMathsEnd[1],
    goodMathsStart[0] + goodMathsMiddle[1] + goodMathsEnd[1],
    goodMathsStart[1] + goodMathsMiddle[0] + goodMathsEnd[0],
    goodMathsStart[1] + goodMathsMiddle[1] + goodMathsEnd[0],
    goodMathsStart[1] + goodMathsMiddle[0] + goodMathsEnd[1],
    goodMathsStart[1] + goodMathsMiddle[1] + goodMathsEnd[1]
)

okMathsStart = (
    "{0} has demonstrated a good knowledge of basic mathematical concepts this year. ",
    "{0} has progressed well in maths this year. ",
    "I have been pleased with {0}’s progress in maths this year. "
)

okMathsEnd = (
    "{3} is able to explain {2} thinking verbally using simple mathematical terms, and is starting to show {2} "
    "working out in {2} book clearly and frequently. {3} can now write complete addition and subtraction equations "
    "in {2} book and has a sound knowledge of number bonds to 10. {4} ability to quickly recall {2} number bonds to "
    "20 is improving, and {1} can correctly write numbers from 0 to 100. {0} demonstrates a good understanding of "
    "the relationship between concrete resources and abstract representations of equations, and can use resources, "
    "such as the Dienes blocks and string-beads to prove {2} answers. I would like to see {0} continue to develop "
    "{2} verbal reasoning and understanding of mathematical terminology. ",
    "{3} can use simple mathematical terms to explain {2} thinking, and is confident about sharing ideas during "
    "maths discussions on the carpet. {0} can correctly write numbers from 0 to 100, and can efficiently recall "
    "{2} number bonds to 10. {3} is developing {2} ability to use {2} number facts to 10 and 20 when solving "
    "addition and subtraction problems, and can now write simple equations clearly in {2} book. {3} is beginning "
    "to show {2} working out in {2} maths book and I would like to see {0} continue to practise explaining the "
    "strategies {1} has used, in order to prove {2} answers and to show an in-depth understanding of concepts. {0} "
    "shows a good comprehension of the ways in which concrete resources, such as the Dienes blocks or string-beads "
    "can be used to represent an abstract equation, and {1} frequently and accurately utilises these resouces when "
    "solving problems during lessons. ",
    "{3} uses simple mathematical terminology to explain {2} thinking and is developing {2} use of more than "
    "one strategy to solve addition and subtraction problems. {0} can quickly recall {2} number bonds to 10, "
    "and is becoming increasingly confident about recalling {2} number facts to 20. I would like to see {0} "
    "begin to apply {2} number facts to 10 and 20 when solving addition and subtraction problems. {0} presents "
    "{2} work neatly in {2} maths book, and {1} can now confidently write simple equations.  {3} is beginning to "
    "demonstrate {2} working out in writing in {2} maths book, and uses concrete resources, such as the Dienes "
    "blocks and string-beads to help {5} solve problems. {0} is increasingly able to show the link between "
    "concrete resources, such as the string-beads or cubes and abstract representations. "
)

badMathsStart = (
    "{0} has made steady progress in maths this year. ",
    "{0} has progressed well in maths this year. "
)

badMathsEnd = (
    "{3} is able to write numbers from 0 to 100 with increasing correctness, and is developing {2} understanding of "
    "more than and less than, as well as {2} ability to order numbers according to size. {3} has developed a sounder "
    "knowledge of number bonds to 10, and is beginning to write complete sums in {2} maths book correctly and with "
    "increasing clarity. I would like {0} to practise recalling {2} number bonds to 10 quickly and correctly, so "
    "that {1} can apply these number facts when solving addition and subtraction facts. {0} is developing a "
    "sounder knowledge of place value, and is increasingly confident when using concrete resources such as the "
    "Dienes blocks and bead-strings to help {5} solve maths problems. ",
    "{3} is able to utilise concrete resources, such as the Dienes blocks and bead-strings to help {5} solve "
    "maths problems, and {1} has gained a sounder knowledge of number bonds to 10. I would like {0} to continue "
    "practising number facts to 10, and to begin consolidating number facts to 20. A good understanding of these "
    "will help {5} when solving more complicated addition and subtraction problems. {0} can confidently and "
    "correctly write the majority of numbers from 0 to 100, and has developed {2} understanding of more than and "
    "less than. {3} can also order numbers according to size with increasing confidence. {0} is developing {2} "
    "knowledge of tens and ones, and is beginning to write complete sums in {2} book correctly. {0} should "
    "practise writing these equations clearly in {2} book. "
)

goodEnglishOverall = (
    "{0} has made excellent progress during English this year. ",
    "{0} has progressed well in English this year. "
)

goodEnglishSentenceStructure = (
    "{3} demonstrates a sound knowledge of a range of sentence structures and can consistently and confidently use "
    "time connectives to write a well structured recount. ",
    "{3} is a confident writer and displays a sound understanding of a variety of sentence structures. "
    "{3} consistently uses a range of time connectives to formulate well-structured recounts and has a good"
    " comprehension of the way in which adjectives and descriptive language can be used to enhance a piece of writing. "
)

goodEnglishPunctuation = (
    "{4} use and comprehension of full stops and capital letters is excellent, and {1} is starting to use other"
    " punctuation, such as speech marks, exclamation marks and possessive apostrophes in {2} writing. ",
    "{4} use of capital letters and full stops is very consistent. "
)

okEnglishProgress = (
    "{0} has made steady progress in English this academic year. ",
    "",
    "{0} has progressed well in English this year. "
)

okEnglishPunctuation = (
    "{3} shows an understanding of the importance of finger spaces as a means to demarcate words in sentences and"
    " {2} use of full stops and capital letters is increasingly accurate and consistent. ",
    "{4} use of finger spaces, capital letters and full stops has improved and {1} shows a good understanding of"
    " the way in which full stops help to demarcate sentences. {0} consistently and accurately distinguishes between"
    " upper and lower case letters, and {2} handwriting is increasingly clear and neat. ",
    "{3} has a sound understanding of capital letters and full stops and uses these in {2} work with increasing"
    " accuracy and consistency. {0} uses finger spaces in order to improve the clarity of {2} writing and {2}"
    " distinction between upper and lower case letters in {2} handwriting has improved considerably. "
)

okEnglishSpelling = (
    "{0} is a confident and imaginative writer and has written a number of creative character descriptions and"
    " stories over the course of the year. {0} is beginning to use time connectives to structure a recount and"
    " {2} correct spelling of high frequency and tricky words has improved considerably since the start of the year. ",
    "{4} accurate spelling of high frequency and tricky words is steadily improving, and {0} should continue to"
    " apply {2} knowledge of phonemes when sounding out and spelling words. {0} is developing {2} knowledge and use of"
    " adjectives, and I have seen {5} use a range of time connectives to write a well-structured recount. ",
    "{0} has a sound knowledge of phase 3 and 4 high frequency words and I would like to see {5} develop more"
    " confidence with spelling phase 5 tricky words. {0} is a confident and imaginative writer and has developed"
    " {2} knowledge of a range of adjectives to enhance {2} writing. "
)

badEnglish = (
    "{0} has become an increasingly confident writer as the year has progressed, and {1} is now able to record "
    "{2} own ideas independently in {2} English book. {0} is always keen to share {2} imaginative ideas during "
    "class discussions, and it has been a delight to see the progress which {1} has made with {2} handwriting and "
    "spelling since the start of the year. {0}'s writing stamina has increased considerably and {1} can now "
    "confidently form full sentences in writing. {4} use of full stops, capital letters and finger spaces is "
    "increasingly consistent. I would like to see {0} continue to practise {2} spellings of high frequency and "
    "tricky words, so that {2} writing is fluent and consistently comprehensible. ",
    "I am pleased with {0}'s steady progress in English, and it has been a pleasure to see {0} grow "
    "in confidence as a writer over the course of the year. {0} is now able to write down {2} own ideas "
    "independently on paper, and {2} writing stamina has improved considerably as the year has progressed. "
    "{3} is able to distinguish between upper and lower case letters and is beginning to use full stops and capital "
    "letters with increasing accuracy. {0} has frequently demonstrated {2} ability to form a full sentence in "
    "writing, and {1} can use 'and' confidently in sentences in order to expand and develop {2} writing. "
    "{3} is beginning to use time connectives to write a structured recount, and {2} correct spelling of high "
    "frequency words is improving steadily. "
)

goodPhonicsOverall = (
    "{0}'s knowledge of phonics sounds is excellent. ",
)

goodPhonicsSounds = (
    "{3} is able to decode and blend a range of words carefully and accurately and can confidently read and transcribe"
    " all of the 40+ phonemes. ",
)

goodPhonicsReading = (
    "{0} shows an enthusiasm for reading and I have been very impressed by {2} insightful and thoughtful"
    " contributions to discussions during  whole-class reading sessions. "
    "{3} demonstrates a good understanding of story features and makes erudite inferences about characters and"
    " events. ",
    "{0} has shown a considerable interest in reading for pleasure, and I have been very impressed by {2} readiness to"
    " share {2} reading experiences and favourite books with the rest of the class. {4} thoughtful contributions to "
    "discussions during whole-class reading sessions have been very much appreciated. "
)

okPhonicsProgress = (
    "{0} has a good knowledge of phonics sounds and can confidently and consistently recognise the majority of"
    " the 40+ phonemes. {4} ability to accurately decode and blend a range of words has improved over the course"
    " of the year, and I would like to see {0} continue to practise sounding out words before writing them down, in"
    " order to improve the accuracy of {2} spelling. ",
    "I am very pleased with {0}'s progress in phonics this year. {3} sounds out and blends a range of words with"
    " increasing confidence and now has a good understanding of a majority of the 40+ phonemes. I would now like"
    " to see {0} develop {2} knowledge of alternative spellings for various sounds. ",
    "{0} is an accurate and confident decoder and blender of words and sounds and I have been pleased with {2}"
    " progress in phonics this year. {3} can consistently identify the majority of the 40+ phonemes and can apply"
    " these sounds well in the majority of cases when spelling words. I would like {0} to take {2} time when"
    " sounding out words for spellings, in order to ensure that {2} spellings are consistently accurate and plausible. "
)

okPhonicsLastSentence = (
    "{0} has become an increasingly confident and accurate reader and I have been impressed by {0}'s readiness"
    " to share {2} ideas about {2} favourite books with the class. ",
    "{0} shows an enthusiasm for reading, and it is a delight to see {0} so engaged during whole-class"
    " reading sessions. {3} makes several insightful and thoughtful contributions to discussions and is frequently"
    " willing to share and discuss {2} favourite books during show and tell. ",
    "{0} is a confident reader, and I am pleased by {2} increasing fluency when reading to an adult. {0} is"
    " frequently keen to share {2} ideas about books with the rest of the class and {1} demonstrates a clear"
    " understanding of the features of a story, such as character and setting. "
)

badPhonics = (
    "{0}'s progress in phonics has been good. {3} is now able to recognise a large number of the 40+ phonemes "
    "with increasing confidence. {0} still finds some sounds tricky, and I would like {0} to continue to "
    "practise {2} phonics sounds as {1} progresses into year two. This will enable {5} to sound out and blend words "
    "more efficiently, and will help to improve {2} spelling skills. {0} has become an increasingly confident and "
    "fluent reader, and {1} frequently shows a readiness to contribute imaginative ideas during whole-class "
    "reading sessions. ",
    "{0} is a confident reader who can sound out and blend a number of words accurately. {4} knowledge of the 40+"
    " phonemes has improved considerably since the start of the year and {1} now has a good knowledge of several"
    " sounds. I would like {0} to continue to practise these sounds as {1} progresses into year two, so that {1} is"
    " able to sound out words with greater speed, and so that {1} can continue to develop the accuracy of {2} "
    "spelling. I am very pleased that {0} is now more confident about contributing to discussions during "
    "whole-class reading sessions. {0} always makes thoughtful and valuable contributions to class discussions "
    "and I would like to see {5} continue to grow in confidence when expressing {2} ideas both verbally and "
    "in writing. "
)

goodScienceStart = (
    "{0} has enjoyed working practically during science lessons this year and has made good progress. {0} has "
    "made insightful and interesting contributions to classroom discussions and {1} shows excellent scientific "
    "reasoning skills. ",
    "{0} is confident in {2} use of scientific vocabulary, and has made interesting and nuanced comments "
    "during classroom discussions. {4} scientific reasoning is excellent and {1} can record {2} findings in "
    "a range of ways. "
)

goodScienceEnd = (
    "{0} has been able to record {2} findings in tables and in writing, and has demonstrated a sound understanding "
    "of the properties of materials, as well as seasons and temperature. ",
    "{3} demonstrates a very good understanding of the properties of materials and showed excellent enquiry "
    "skills during an investigation about opaque and transparent materials. "
)

okScience = (
    "{0} has enjoyed taking part in a variety of scientific experiments this year and has developed {2} "
    "reasoning skills. {3} is beginning to record {2} findings in a range of ways and has been able to identify "
    "various materials and some of their properties. {0} is beginning to develop {2} use of scientific "
    "vocabulary, and has made successful comparisons between humans, plants and animals. ",
    "{0} has shown an interest and curiosity to know more in science this year and {1} is beginning to use "
    "scientific language more confidently and accurately. When carrying out experiments, {0} is able to make "
    "sensible predictions and conclusions. {3} is beginning to record {2} ideas in charts and in writing, and has "
    "shown a good understanding of the properties of a range of materials. ",
    "{0} has expressed enthusiasm and intrigue when taking part in a range of scientific experiments this "
    "year. {3} is starting to record {2} findings in a variety of different forms, and is beginning to use "
    "scientific vocabulary to explain {2} ideas. I have been impressed by {0}’s thoughtful contributions to "
    "class discussions in science lessons and {1} can indentify a range of materials and their properties. "
)

badScienceStart = (
    "{0} has enjoyed working practically during science this year and has shown an interest in the process of "
    "scientific investigation. ",
    "{0} has shown an interest in science this year, and has worked well with {2} peers during "
    "scientific investigations. "
)

badScienceEnd = (
    "{3} is good at collaborating with {2} peers and has been able to make simple predictions. {3} contributes "
    "increasingly frequently during class discussions and can identify and describe a range of materials. ",
    "{3} is beginning to use simple scientific language to explain {2} ideas, and successfully identified some "
    "opaque, transparent and translucent materials during an experiment about light and shadow. "
)

goodReligStudyStart = (
    "{0} has made excellent progress during R.E this year. ",
    "{0} has progressed very well in R.E this year. "
)

goodReligStudyEnd = (
    "{3} has shown a high level of respect when listening to Bible stories and has been able to draw meaning "
    "from them, as well as retell these stories in {2} own words, using a range of religious terminology. {3} has "
    "made frequent, insightful contributions to class discussions on a range of topics, such as the role of the "
    "church in the religious community. {3} can eloquently explain the meaning of important religious festivals, "
    "such as Christmas, Easter and Pentecost. {0} engages in class and assembly worship in a consistently "
    "prayerful and respectful manner, and demonstrates a good knowledge of the role which God plays "
    "in our own lives. ",
    "{3} acts in a prayerful and mature manner during times of collective worship, and is able to recall and "
    "lead prayers confidently and accurately. {3} listens respectfully and attentively to religious stories "
    "during R.E lessons, and demonstrates a sound understanding of the link between the teachings of the Bible "
    "and the ways in which God helps us in our own lives. {0} can explain the meaning of these stories both "
    "verbally and in writing, and consistently makes a range of insightful and thoughtful contributions to class "
    "discussions. {3} has shown an excellent understanding of the meaning of important religious festivals, such as "
    "Christmas, Easter and Pentecost. "
)

okReligStudy = (
    "{0} has made good progress in R.E this year. {3} has listened to religious stories with increasing "
    "attention throughout the year and always shows a high level of respect during times of collective worship "
    "in the classroom and in assemblies. {0} is able to explain the meaning of Biblical stories in {2} own "
    "words with increasing confidence and has developed {2} understanding of the role which God plays in our own "
    "lives. {0} can recall important prayers well and participates enthusiastically during class discussions. "
    "{3} demonstrates a good understanding of the meaning of important religious festivals, such as Christmas, "
    "Easter and Pentecost. ",
    "{0} always acts prayerfully and respectfully during times of collective worship and has engaged well "
    "during class discussions on a range of topics, such as the role of the church within the community. {3} is "
    "able to draw meaning from religious stories during lessons and can retell these stories in {2} own words "
    "using simple religious vocabulary. {3} shows a good knowledge of important religious festivals, such as "
    "Christmas, Easter and Pentecost, and can explain their significance. {0} is increasingly able to make "
    "the link between the teachings of the Bible and the way in which God helps us in our own lives. ",
    "{0} has enjoyed listening to a variety of Biblical stories over the course of the year and {1} responds "
    "to them well by frequently contributing to class discussions on a range of topics. {3} is beginning to "
    "understand the messages which these stories convey and can retell their events using simple religious "
    "vocabulary. {0} has written some good recounts of important religious festivals, such as Christmas, Easter "
    "and Pentecost and can explain their meaning and significance. {0} always acts respectfully and prayerfully "
    "during times of collective worship, and is beginning to shown an understanding of how God helps us in "
    "our own lives. "
)

badReligStudyStart = (
    "{0} has shown maturity and respect during times of collective worship both in the classroom and in "
    "assemblies. {3} joins in confidently with important prayers and hymns and listens attentively to religious "
    "stories. {3} is beginning to retell some of these stories in {2} own words and recently demonstrated a good "
    "understanding of the story of Pentecost by sequencing the events correctly. ",
    "{0} has listened to religious stories during R.E lessons with increasing attention throughout the year and "
    "is beginning to draw meaning from them. {3} has become more confident about retelling the main events of these "
    "stories, and can identify important religious festivals, such as Christmas, Easter and Pentecost. "
)

badReligStudyEnd = (
    "{0} has begun to contribute more frequently to class discussions, and I would like to see {0} develop "
    "{2} confidence in sharing {2} ideas with {2} peers, as {1} has lots of thoughtful and interesting ideas to "
    "share. {0} is beginning to show an understanding the ways in which God helps us in our own lives, and {1} "
    "lives out Christian values every day by showing {2} peers kindness and respect. ",
    "{3} is beginning to explain the significance and meaning of the events and is starting to link the meaning of "
    "these stories, to the way in which God helps us in own lives. {3} shows a good understanding of Christians "
    "values and treats {2} peers with kindness and respect. "
)

goodOtherSubjects = (
    "{0} has worked very well across all areas of the curriculum. {3} shows an enthusiasm and interest "
    "in many areas, such as ",
    "{0} has applied {5}self well across all areas of the curriculum. {3} has worked particularly well in "
)

report += "\nRELIGIOUS STUDIES: \n"
if religStudy == 1:
    x = random.randint(0, len(badReligStudyStart) - 1)
    religStudyReport += badReligStudyStart[x]
    x = random.randint(0, len(badReligStudyEnd) - 1)
    religStudyReport += badReligStudyEnd[x]
elif religStudy == 2:
    x = random.randint(0, len(okReligStudy) - 1)
    religStudyReport += okReligStudy[x]
else:
    x = random.randint(0, len(goodReligStudyStart) - 1)
    religStudyReport += goodReligStudyStart[x]
    x = random.randint(0, len(goodReligStudyEnd) - 1)
    religStudyReport += goodReligStudyEnd[x]

report += religStudyReport + "\n"

report += "\nENGLISH: \n"
if english == 1:
    x = random.randint(0, len(badEnglish) - 1)
    englishReport += badEnglish[x]
elif english == 2:
    x = random.randint(0, len(okEnglishProgress) - 1)
    englishReport += okEnglishProgress[x]
    x = random.randint(0, len(okEnglishPunctuation) - 1)
    englishReport += okEnglishPunctuation[x]
    x = random.randint(0, len(okEnglishSpelling) - 1)
    englishReport += okEnglishSpelling[x]
else:
    x = random.randint(0, len(goodEnglishOverall) - 1)
    englishReport += goodEnglishOverall[x]
    x = random.randint(0, len(goodEnglishSentenceStructure) - 1)
    englishReport += goodEnglishSentenceStructure[x]
    x = random.randint(0, len(goodEnglishPunctuation) - 1)
    englishReport += goodEnglishPunctuation[x]

englishReport += "\n\n"

if phonics == 1:
    x = random.randint(0, len(badPhonics) - 1)
    englishReport += badPhonics[x]
elif phonics == 2:
    x = random.randint(0, len(okPhonicsProgress) - 1)
    englishReport += okPhonicsProgress[x]
    x = random.randint(0, len(okPhonicsLastSentence) - 1)
    englishReport += okPhonicsLastSentence[x]
else:
    x = random.randint(0, len(goodPhonicsOverall) - 1)
    englishReport += goodPhonicsOverall[x]
    x = random.randint(0, len(goodPhonicsSounds) - 1)
    englishReport += goodPhonicsSounds[x]
    x = random.randint(0, len(goodPhonicsReading) - 1)
    englishReport += goodPhonicsReading[x]

report += englishReport + "\n"

report += "\nMATHS: \n"
if maths == 1:
    x = random.randint(0, len(badMathsStart) - 1)
    mathsReport += badMathsStart[x]
    x = random.randint(0, len(badMathsEnd) - 1)
    mathsReport += badMathsEnd[x]
elif maths == 2:
    x = random.randint(0, len(okMathsStart) - 1)
    mathsReport += okMathsStart[x]
    x = random.randint(0, len(okMathsEnd) - 1)
    mathsReport += okMathsEnd[x]
else:
    x = random.randint(0, len(goodMaths) - 1)
    mathsReport += goodMaths[x]

report += mathsReport + "\n"

report += "\nSCIENCE: \n"
if science == 1:
    x = random.randint(0, len(badScienceStart) - 1)
    scienceReport += badScienceStart[x]
    x = random.randint(0, len(badScienceEnd) - 1)
    scienceReport += badScienceEnd[x]
elif science == 2:
    x = random.randint(0, len(okScience) - 1)
    scienceReport += okScience[x]
else:
    x = random.randint(0, len(goodScienceStart) - 1)
    scienceReport += goodScienceStart[x]
    x = random.randint(0, len(goodScienceEnd) - 1)
    scienceReport += goodScienceEnd[x]

report += scienceReport + "\n"


report += "\nGENERAL COMMENTS: \n"
x = random.randint(0, len(goodOtherSubjects) - 1)
generalReport += goodOtherSubjects[x]

report += generalReport + "\n"

religStudyReport = religStudyReport.format(firstName,
                       nominalPronoun,
                       possessivePronoun,
                       nominalPronounCapitalised,
                       possessivePronounCapitalised,
                       pointingPronoun,
                       pointingPronounCapitalised)
englishReport = englishReport.format(firstName,
                       nominalPronoun,
                       possessivePronoun,
                       nominalPronounCapitalised,
                       possessivePronounCapitalised,
                       pointingPronoun,
                       pointingPronounCapitalised)
mathsReport = mathsReport.format(firstName,
                       nominalPronoun,
                       possessivePronoun,
                       nominalPronounCapitalised,
                       possessivePronounCapitalised,
                       pointingPronoun,
                       pointingPronounCapitalised)
scienceReport = scienceReport.format(firstName,
                       nominalPronoun,
                       possessivePronoun,
                       nominalPronounCapitalised,
                       possessivePronounCapitalised,
                       pointingPronoun,
                       pointingPronounCapitalised)
generalReport = generalReport.format(firstName,
                       nominalPronoun,
                       possessivePronoun,
                       nominalPronounCapitalised,
                       possessivePronounCapitalised,
                       pointingPronoun,
                       pointingPronounCapitalised)
report = report.format(firstName,
                       nominalPronoun,
                       possessivePronoun,
                       nominalPronounCapitalised,
                       possessivePronounCapitalised,
                       pointingPronoun,
                       pointingPronounCapitalised)


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def addRating(rating):
    if rating == 1:
        return "WT"
    elif rating == 2:
        return "At"
    else:
        return "Ab"


def addDetails():
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if 'Pupil:' in paragraph.text:
                        paragraph.add_run(" " + fullName, style = 'NameStyle')
                    if 'Class:' in paragraph.text:
                        paragraph.add_run(" Year 1", style = 'NameStyle')
                    if 'Teacher:' in paragraph.text:
                        paragraph.add_run(" Miss B. Archibald", style = 'NameStyle')
                        return


def addSubjectAttainment(subject, subjectAttainment):
    for table in document.tables:
        lstRow = iter(table.rows)
        for row in lstRow:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if subject in paragraph.text:
                        for rowExtra in lstRow:
                            for cellExtra in rowExtra.cells:
                                for paraExtra in cellExtra.paragraphs:
                                    if 'Achievement' in paraExtra.text:
                                        for rowExtraExtra in lstRow:
                                            for cellExtraExtra in rowExtraExtra.cells:
                                                for paraExtraExtra in cellExtraExtra.paragraphs:
                                                    if paraExtraExtra.text == "":
                                                        paraExtraExtra.add_run(addRating(subjectAttainment),
                                                                               style = 'ReportStyle')
                                                        return


def addSubjectEffort(subject, subjectEffort):
    for table in document.tables:
        lstColumn = iter(table.columns)
        lstRow = iter(table.rows)
        for row in lstRow:
            for column in lstColumn:
                for cell in column.cells:
                    for paragraph in cell.paragraphs:
                        if subject in paragraph.text:
                            for colExtra in lstColumn:
                                for cellExtra in colExtra.cells:
                                    for paraExtra in cellExtra.paragraphs:
                                        if 'Effort' in paraExtra.text:
                                            for rowExtra in lstRow:
                                                for cellExtraExtra in rowExtra.cells:
                                                    for paraExtraExtra in cellExtraExtra.paragraphs:
                                                        if paraExtraExtra.text == "":
                                                            paraExtraExtra.add_run(addRating(subjectEffort),
                                                                                   style = 'ReportStyle')
                                                            return


def addOtherSubjectsReport(otherReport):
    step = 0
    for table in document.tables:
        lstColumn = iter(table.columns)
        lstColumn2 = iter(table.columns)
        for column in lstColumn:
            for cell in column.cells:
                for paragraph in cell.paragraphs:
                    if 'Other Subjects' in paragraph.text:
                        lstRow = iter(table.rows)
                        for columnExtra in lstColumn:
                            for cellExtra in columnExtra.cells:
                                for paraExtra in cellExtra.paragraphs:
                                    if 'Comment' in paraExtra.text:
                                        next(lstRow)
                                        for row in lstRow:
                                            for rowCell in row.cells:
                                                for rowCellPara in rowCell.paragraphs:
                                                    if rowCellPara.text == "" and step == 2:
                                                        rowCellPara.add_run(otherReport, style = 'ReportStyle')
                                                        return
                                                    elif rowCellPara.text == "" and step != 2:
                                                        step += 1
                                                    
                                                    
def addSubjectReportText(subject, subjectReport):
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                lst = iter(cell.paragraphs)
                for paragraph in lst:
                    if subject in paragraph.text:
                        paragraph.add_run("\n" + subjectReport, style = 'ReportStyle')
                        for para in lst:
                            delete_paragraph(para)
                        return


def addSubjectReport(subject, subjectReport, subjectAttainment, subjectEffort):
    addSubjectReportText(subject, subjectReport)
    addSubjectEffort(subject, subjectEffort)
    addSubjectAttainment(subject, subjectAttainment)
    return


addDetails()
addSubjectReport('Religious Education', religStudyReport, religStudy, religStudyEffort)
addSubjectReport('English', englishReport, (english + phonics) // 2, (englishEffort + phonicsEffort) // 2)
addSubjectReport('Mathematics', mathsReport, maths, mathsEffort)
addSubjectReport('Science', scienceReport, science, scienceEffort)
addOtherSubjectsReport(generalReport)

documentName = firstName + ' ' + lastName + '.docx'
document.save(documentName)

print(report)
