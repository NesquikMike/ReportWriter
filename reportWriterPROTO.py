import sys
import random
from docx import Document

document = Document('master.docx')

if len(sys.argv) < 5:
    print("Not enough information. Please provide First Name, Last Name, Gender - f for female, m for male,"
          " Attainment and Behaviour - where 1 is bad, 2 is ok, 3 is good attainment or behaviour.")
    exit()

firstName = sys.argv[1]
secondName = sys.argv[2]
fullName = firstName + " " + secondName
gender = sys.argv[3]
attainment = int(sys.argv[4])
behaviour = int(sys.argv[5])
nominalPronoun = "she"
possessivePronoun = "her"
report = ""

if gender != "m" and gender != "f":
    print("Sorry, this gender is not yet supported. Please enter f for female or m for male")
    exit()
elif gender == "m":
    nominalPronoun = "he"
    possessivePronoun = 'his'

if attainment < 1 or attainment > 3:
    print("Please insert a valid value for Attainment- where 1 is bad, 2 is ok, 3 is good Attainment.")
    exit()

if behaviour < 1 or behaviour > 3:
    print("Please insert a valid value for Attainment- where 1 is bad, 2 is ok, 3 is good Behaviour.")
    exit()

goodAttainment = (
    "{0} {1} has made excellent progress this academic year, {2} has made fantastic progress in Phonics especially. ",
    "{0} has done very well this year, {2} has shown fantastic ability in Maths. ",
    "{0} {1} has shown fantastic ability this year, {2} has made progress in all of"
    " {2} subjects and is in a great place to move into Year 2. "
)

okAttainment = (
    "{0} {1} has made reasonable progress this academic year, {2} has made the most progress in Science. ",
    "{0} has done well this year, {2} should focus on {2} phonics where more improvement is needed. ",
    "{0} {1} has shown good ability this year, {2} should practice Maths more"
    " in order to be best placed to move into Year 2. "
)

badAttainment = (
    "{0} {1} has not made enough progress this academic year, {2} needs to improve in English to be ready for the"
    " jump to Year 2. ",
    "{0} has not done well enough this year, {2} has shown glimmers of ability in Maths,"
    " but needs to improve in Phonics. ",
    "{0} {1} has not shown enough ability this year, {2} has not made progress in all of"
    " their subjects and should be worried about moving into Year 2. "
)

goodBehaviour = (
    "{0} is very well behaved and sets a fantastic example to {2} peers. ",
    "{0} is always ready to learn and leads the classroom in behaviour. ",
    "{0} is delightful to work with and always gives 100%. "
)

okBehaviour = (
    "{0} generally behaves well, but is occasionally distracted by classmates. ",
    "{0} can be distracted by others, but generally is studious and applies a good work ethic. ",
    "{0}'s behaviour is inconsistent, {2} at times sets a fantastic example while at other times is very disruptive. "
)

badBehaviour = (
    "{0} is too chatty sometimes and would do well not to stay focused on their work instead of distracting others. ",
    "{0}'s energy needs to be applied more to studies and less to playing with classmates. ",
    "{0} is a disruptive influence in the class and would benefit from focusing more in lessons. "
)
praise = (
    "Great work {0}!",
    "Excellent {0}!",
    "Superb {0}!",
    "Keep up the excellent work {0}!"
)

decent = (
    "Keep it up {0}!",
    "Push it to next level {0}!",
    "Work harder to show more of your ability {0}!"
)

if attainment == 1:
    x = random.randint(0, len(badAttainment) - 1)
    report += badAttainment[x].format(firstName, secondName, nominalPronoun)
elif attainment == 2:
    x = random.randint(0, len(okAttainment) - 1)
    report += okAttainment[x].format(firstName, secondName, nominalPronoun)
else:
    x = random.randint(0, len(goodAttainment) - 1)
    report += goodAttainment[x].format(firstName, secondName, nominalPronoun)

#report += "\n"

if behaviour == 1:
    x = random.randint(0, len(badBehaviour) - 1)
    report += badBehaviour[x].format(firstName, secondName, nominalPronoun)
elif behaviour == 2:
    x = random.randint(0, len(okBehaviour) - 1)
    report += okBehaviour[x].format(firstName, secondName, nominalPronoun)
else:
    x = random.randint(0, len(goodBehaviour) - 1)
    report += goodBehaviour[x].format(firstName, secondName, nominalPronoun)

#report += "\n"

if (attainment + behaviour) > 4:
    x = random.randint(0, len(praise) - 1)
    report += praise[x].format(firstName, secondName, nominalPronoun)
else:
    x = random.randint(0, len(decent) - 1)
    report += decent[x].format(firstName, secondName, nominalPronoun)

print(report)

def addText():
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if 'Pupil:' in paragraph.text:
                        paragraph.add_run(" " + fullName)
                    if 'Class:' in paragraph.text:
                        paragraph.add_run(" Year 1")
                    if 'Teacher:' in paragraph.text:
                        paragraph.add_run(" Miss B. Archibald")
                    if 'Religious Education' in paragraph.text:
                        paragraph.add_run("\n" + report)
                        return


addText()

#document.add_paragraph(report)
documentName = firstName + secondName + '.docx'
document.save(documentName)
# Maths English Phonics Science R.E. Behavioural/Application(Optional) for each subject General comment about behaviour